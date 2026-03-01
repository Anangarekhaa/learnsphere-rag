from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Body,Depends, HTTPException
import shutil
import os
from fastapi import Security
from pydantic import BaseModel, EmailStr, Field
from fastapi.security import OAuth2PasswordBearer
from sqlalchemy.orm import Session, relationship
from requests import Session
from app.questionnaire_service import process_questionnaire
from app.database import SessionLocal,engine,Base
from app.models import *
from app.questionnaire_parser import parse_questionnaire
from app.qa_service import answer_question
from uuid import UUID
from fastapi.responses import FileResponse
from docx import Document as DocxDocument
from fastapi.middleware.cors import CORSMiddleware
from typing import List
from jose import jwt, JWTError
from fastapi.security import OAuth2PasswordBearer
from app.database import get_db
from app.auth import get_current_user
from app.auth import hash_password, verify_password, create_access_token, SECRET_KEY, ALGORITHM
from app.chunking import chunk_text
from app.embeddings import get_embedding


load_dotenv()

Base.metadata.create_all(bind=engine)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)



UPLOAD_FOLDER = "uploaded_files"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

EXPORT_FOLDER = "exports"
os.makedirs(EXPORT_FOLDER, exist_ok=True)



@app.get("/")
def root():
    return {"message": "LearnSphere Questionnaire Tool Running"}


@app.post("/upload-questionnaire")
def upload_questionnaire(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
   

    unique_filename = f"{uuid.uuid4()}_{file.filename}"
    file_path = os.path.join(UPLOAD_FOLDER, unique_filename)

   
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

   
    run = QuestionnaireRun(
        user_id=current_user.id,
        filename=file.filename
    )
    db.add(run)
    db.commit()
    db.refresh(run)

    
    questions = parse_questionnaire(file_path)

    results = []
    answered = 0
    total_confidence = 0

    confidence_map = {
    "High": 3,
    "Medium": 2,
    "Low": 1
    }


    for q in questions:
        response = answer_question(q,current_user.id)

        qa = QuestionAnswer(
            run_id=run.id,
            question=q,
            answer=response["answer"],
            confidence=response["confidence"],
            citation=response["citations"][0] if response["citations"] else None
        )

        db.add(qa)
        db.flush()
        
        if response["confidence"] != "None":
            answered += 1
            total_confidence += confidence_map.get(response["confidence"], 0)

        results.append({
            "question_id": str(qa.id),
            "question": q,
            "answer": response["answer"],
            "confidence": response["confidence"],
            "citation": response["citations"][0] if response["citations"] else None
        })

    db.commit()

    summary = {
        "total_questions": len(questions),
        "answered_with_citations": answered,
        "not_found": len(questions) - answered,
        "coverage_percent": round((answered / len(questions)) * 100),
        "average_confidence": round(total_confidence / answered) if answered else 0
    }

    return {
        "run_id": str(run.id),
        "filename": file.filename,
        "summary": summary,
        "results": results
    }


@app.get("/results/{run_id}")
def get_results(run_id: UUID):
    db = SessionLocal()

    run = db.query(QuestionnaireRun).filter_by(id=run_id).first()

    if not run:
        db.close()
        return {"error": "Run not found"}

    answers = (
        db.query(QuestionAnswer)
        .filter_by(run_id=run_id)
        .order_by(QuestionAnswer.id)
        .all()
    )

    total_questions = len(answers)

    answered_with_citation = sum(
        1 for a in answers
        if a.answer != "Not found in references." and a.citation
    )

    not_found_count = sum(
        1 for a in answers
        if a.answer == "Not found in references."
    )

    avg_confidence = (
        round(sum(a.confidence for a in answers) / total_questions)
        if total_questions > 0 else 0
    )
    coverage_percent = round((answered_with_citation / total_questions) * 100) if total_questions > 0 else 0

    response = {
        "run_id": str(run.id),
        "filename": run.filename,
        "summary": {
            "total_questions": total_questions,
            "answered_with_citations": answered_with_citation,
            "not_found": not_found_count,
            "average_confidence": avg_confidence,
            "coverage_percent": coverage_percent
        },
        "results": [
            {
                "question_id": str(a.id),
                "question": a.question,
                "answer": a.answer,
                "confidence": a.confidence,
                "citation": a.citation,
            }
            for a in answers
        ],
    }

    db.close()
    return response


@app.put("/answers/{question_id}")
def update_answer(question_id: UUID, updated_answer: str = Body(...)):
    db = SessionLocal()

    qa = db.query(QuestionAnswer).filter_by(id=question_id).first()

    if not qa:
        db.close()
        return {"error": "Answer not found"}

    qa.answer = updated_answer
    db.commit()

    db.close()

    return {
        "message": "Answer updated successfully",
        "question_id": str(question_id)
    }



@app.get("/export/{run_id}")
def export(
    run_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    run = db.query(QuestionnaireRun).filter(
        QuestionnaireRun.id == run_id,
        QuestionnaireRun.user_id == current_user.id
    ).first()

    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    answers = db.query(QuestionAnswer).filter(
        QuestionAnswer.run_id == run_id
    ).all()

    doc = DocxDocument()

    for qa in answers:
        doc.add_paragraph(f"Question: {qa.question}")
        doc.add_paragraph(f"Answer: {qa.answer}")
        doc.add_paragraph("")

    file_path = os.path.join(EXPORT_FOLDER, f"{run_id}.docx")
    doc.save(file_path)

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="questionnaire_response.docx"
    )




class UserCreate(BaseModel):
    email: EmailStr
    password: str = Field(min_length=8, max_length=64)


@app.post("/signup")
def signup(user: UserCreate, db: Session = Depends(get_db)):
    existing = db.query(User).filter(User.email == user.email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    new_user = User(
        email=user.email,
        hashed_password=hash_password(user.password)
    )

    db.add(new_user)
    db.commit()

    return {"message": "User created successfully"}


class LoginRequest(BaseModel):
    email: EmailStr
    password: str = Field(min_length=8, max_length=64)
    
@app.post("/login")
def login(data: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == data.email).first()

    if not user or not verify_password(data.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")

    token = create_access_token({"sub": user.email})

    return {
        "access_token": token,
        "token_type": "bearer"
    }

class UpdatedAnswer(BaseModel):
    question_id: str
    answer: str
@app.post("/update-answers/{run_id}")
def update_answers(
    run_id: str,
    updated_answers: List[UpdatedAnswer],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    for item in updated_answers:
        qa = db.query(QuestionAnswer).filter(
            QuestionAnswer.id == item.question_id,
            QuestionAnswer.run_id == run_id
        ).first()

        if qa:
            qa.answer = item.answer

    db.commit()

    return {"message": "Answers updated successfully"}


@app.post("/upload-reference")
def upload_reference_document(
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    for file in files:
        content = file.file.read().decode("utf-8")

        document = Document(
            filename=file.filename,
            content=content,
            user_id=current_user.id
        )

        db.add(document)
        db.flush()

        chunks = chunk_text(content)

        for chunk in chunks:
            embedding = get_embedding(chunk)

            db_chunk = DocumentChunk(
                document_id=document.id,
                chunk_text=chunk,
                embedding=embedding
            )

            db.add(db_chunk)

    db.commit()

    return {"message": "Reference documents uploaded successfully"}


